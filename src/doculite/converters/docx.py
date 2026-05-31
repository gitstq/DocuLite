"""
📝 DocuLite DOCX Converter Module

Converter for Microsoft Word documents.
"""

from pathlib import Path
from typing import Optional, Dict, Any, Union, List
import io
import re

from .base import BaseConverter
from ..types import ConversionResult, DocumentType
from ..exceptions import ConversionError
from ..utils import clean_text


class DocxConverter(BaseConverter):
    """Converter for Microsoft Word documents (.docx, .doc)."""
    
    SUPPORTED_TYPES = [DocumentType.DOCX, DocumentType.DOC]
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__(config)
        
    def _get_docx_module(self):
        """Lazy load python-docx."""
        try:
            import docx
            return docx
        except ImportError:
            raise ConversionError(
                "python-docx is required for Word document conversion. "
                "Install with: pip install 'doculite[docx]'"
            )
    
    def convert(self, file_path: Union[str, Path, io.BytesIO], **kwargs) -> ConversionResult:
        """
        Convert a Word document to markdown.
        
        Args:
            file_path: Path to the document or file-like object
            **kwargs: Additional options
            
        Returns:
            ConversionResult with markdown content
        """
        docx = self._get_docx_module()
        
        try:
            if isinstance(file_path, (str, Path)):
                path = self.validate_file(file_path)
                doc = docx.Document(str(path))
            else:
                file_path.seek(0)
                doc = docx.Document(file_path)
            
            markdown_parts = []
            metadata = {}
            pages_info = []
            images_info = []
            tables_info = []
            links_info = []
            
            # Extract core properties
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "revision": core_props.revision or 0,
            }
            
            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    md_text = self._convert_paragraph(para)
                    if md_text:
                        markdown_parts.append(md_text)
            
            # Process tables
            for table_idx, table in enumerate(doc.tables, start=1):
                table_md = self._convert_table(table)
                if table_md:
                    tables_info.append({
                        "index": table_idx,
                        "rows": len(table.rows),
                        "columns": len(table.columns),
                    })
                    markdown_parts.append(table_md)
            
            # Combine all parts
            markdown = "\n\n".join(markdown_parts)
            markdown = self.post_process_markdown(markdown)
            
            # Generate plain text
            text = re.sub(r'[#*_`\[\]|]', '', markdown)
            text = clean_text(text)
            
            return ConversionResult(
                markdown=markdown,
                text=text,
                metadata=metadata,
                pages=pages_info,
                images=images_info,
                tables=tables_info,
                links=links_info,
            )
            
        except Exception as e:
            raise ConversionError(
                f"Failed to convert Word document: {str(e)}",
                {"file_path": str(file_path) if isinstance(file_path, (str, Path)) else "stream"}
            )
    
    def _convert_paragraph(self, para) -> str:
        """Convert a paragraph to markdown."""
        text = para.text.strip()
        if not text:
            return ""
        
        # Check paragraph style for heading
        style_name = para.style.name.lower() if para.style else ""
        
        # Handle headings
        if 'heading' in style_name:
            level = self._extract_heading_level(style_name)
            prefix = '#' * level
            return f"{prefix} {text}"
        
        # Handle list items
        if para._p is not None:
            # Check if it's a list item
            num_pr = para._p.pPr.numPr if para._p.pPr is not None else None
            if num_pr is not None:
                # It's a numbered list
                return f"1. {text}"
            
            # Check for bullet points
            if style_name.startswith('list'):
                return f"- {text}"
        
        # Handle inline formatting
        text = self._process_inline_formatting(para)
        
        return text
    
    def _extract_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        match = re.search(r'heading\s*(\d)', style_name.lower())
        if match:
            return min(int(match.group(1)), 6)
        return 2  # Default to level 2
    
    def _process_inline_formatting(self, para) -> str:
        """Process inline formatting (bold, italic, etc.)."""
        parts = []
        
        for run in para.runs:
            text = run.text
            if not text:
                continue
            
            # Apply formatting
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            
            if run.underline:
                text = f"<u>{text}</u>"
            
            # Handle code/monospace
            if run.font and run.font.name and 'mono' in run.font.name.lower():
                text = f"`{text}`"
            
            parts.append(text)
        
        return ''.join(parts)
    
    def _convert_table(self, table) -> str:
        """Convert a table to markdown."""
        if not table.rows:
            return ""
        
        rows = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            rows.append(row_cells)
        
        if not rows:
            return ""
        
        # Build markdown table
        lines = []
        
        # Header row
        header = rows[0]
        lines.append("| " + " | ".join(header) + " |")
        
        # Separator
        separator = "| " + " | ".join(["---"] * len(header)) + " |"
        lines.append(separator)
        
        # Data rows
        for row in rows[1:]:
            # Ensure row has same number of columns as header
            while len(row) < len(header):
                row.append("")
            row = row[:len(header)]  # Truncate if too long
            lines.append("| " + " | ".join(row) + " |")
        
        return '\n'.join(lines)
